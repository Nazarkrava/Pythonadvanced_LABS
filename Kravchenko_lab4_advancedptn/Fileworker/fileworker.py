import sqlite3

import pandas as pd
from docx import Document

from Utils.settings import DB_PATH, DOC_PATH, XLSX_PATH


class Fileworker:
    def __init__(self):
        print('Saving')

    def select_data(self):
        self.cars = pd.read_sql("select * from clients", sqlite3.connect(DB_PATH))
        self.types_of_violantions = pd.read_sql("select * from directions", sqlite3.connect(DB_PATH))
        self.acts_of_violations = pd.read_sql("select * from trips", sqlite3.connect(DB_PATH))

    def delete_files(self):
        import os
        files = [
                 XLSX_PATH + "Clients.xlsx",
                 XLSX_PATH + "Directions.xlsx",
                 XLSX_PATH + "Trips.xlsx",
                 DOC_PATH + 'Clients.docx',
                 DOC_PATH + 'Directions.docx',
                 DOC_PATH + 'Trips.docx'
                 ]
        for file_path in files:
            if os.path.exists(file_path):
                os.remove(file_path)

    def save_to_files(self):
        self.delete_files()
        self.select_data()
        self.save_to_xlsx()
        self.save_to_doc()
        print('Saved')
    def save_to_xlsx(self):
        self.cars.to_excel(XLSX_PATH + "Clients.xlsx")
        self.types_of_violantions.to_excel(XLSX_PATH + "Directions.xlsx")
        self.acts_of_violations.to_excel(XLSX_PATH + "Trips.xlsx")

    def save_to_doc(self):
        self._save_to_doc(self.cars, 'cars', DOC_PATH + 'Clients.docx')
        self._save_to_doc(self.types_of_violantions, 'Directions', DOC_PATH + 'Directions.docx')
        self._save_to_doc(self.acts_of_violations, 'Trips', DOC_PATH + 'Trips.docx')

    def _save_to_doc(self, data, tabletitle, path_to_file):
        document = Document()
        data = pd.DataFrame(data)
        document.add_heading(tabletitle)
        table = document.add_table(rows=(data.shape[0]), cols=data.shape[1])
        for i, column in enumerate(data):
            for row in range(data.shape[0]):
                table.cell(row, i).text = str(data[column][row])
        document.save(path_to_file)
