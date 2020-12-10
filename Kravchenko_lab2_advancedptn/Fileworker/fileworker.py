import sqlite3

import pandas as pd
from docx import Document

from Utils.settings import DB_PATH, DOC_PATH, XLSX_PATH


class Fileworker:
    def __init__(self):
        print()

    def select_data(self):
        self.clients = pd.read_sql("select * from clients", sqlite3.connect(DB_PATH))
        self.types_of_violantions = pd.read_sql("select * from directions", sqlite3.connect(DB_PATH))
        self.acts_of_violations = pd.read_sql("select * from trips", sqlite3.connect(DB_PATH))

    def delete_files(self):
        import os
        files = [
                XLSX_PATH + "clients.xlsx",
                 XLSX_PATH + "directions.xlsx",
                 XLSX_PATH + "acts_of_violations.xlsx",
                 DOC_PATH + 'clients.docx',
                 DOC_PATH + 'directions.docx',
                 DOC_PATH + 'trips.docx'
                 ]
        for file_path in files:
            if os.path.exists(file_path):
                os.remove(file_path)

    def save_to_files(self):
        self.delete_files()
        self.select_data()
        self.save_to_xlsx()
        self.save_to_doc()

    def save_to_xlsx(self):
        self.clients.to_excel(XLSX_PATH + "clients.xlsx")
        self.types_of_violantions.to_excel(XLSX_PATH + "directions.xlsx")
        self.acts_of_violations.to_excel(XLSX_PATH + "trips.xlsx")

    def save_to_doc(self):
        self._save_to_doc(self.clients, 'clients', DOC_PATH + 'clients.docx')
        self._save_to_doc(self.types_of_violantions, 'directions', DOC_PATH + 'directions.docx')
        self._save_to_doc(self.acts_of_violations, 'trips', DOC_PATH + 'trips.docx')

    def _save_to_doc(self, data, tabletitle, path_to_file):
        document = Document()
        data = pd.DataFrame(data)
        document.add_heading(tabletitle)
        table = document.add_table(rows=(data.shape[0]), cols=data.shape[1])
        for i, column in enumerate(data):
            for row in range(data.shape[0]):
                table.cell(row, i).text = str(data[column][row])
        document.save(path_to_file)
